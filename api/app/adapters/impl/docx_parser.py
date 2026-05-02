"""python-docx 기반 DOCX 문서 파서 (W5 DE-67).

배경
- W4-Q-9 sniff (`work-log/2026-05-02 W4-Q-9 sniff DOCX·PPTX 라이브러리 평가.md`) 결과
  HwpxParser 패턴 (sticky propagate + style 기반 heading) 직접 재사용 가능 판정.
- 페르소나 A 의 회의록·보고서·메모 자료 빈도 ↑ → DE-67 (a) heuristic-only 채택.

설계
- `iter_inner_content()` (python-docx 1.x) 로 paragraph + table 의 XML 순서 순회
  → DOCX 의 본문/표 인터리빙 보존
- heading 판별: (A) `paragraph.style.name` 의 'Heading' 패턴 → (B) inline 텍스트 패턴 fallback
  (HwpxParser 와 동일 알고리즘)
- 표 처리: 각 표는 별도 `ExtractedSection` (행/셀 텍스트를 ` | ` 로 join) — chunk_filter
  의 table_noise 룰이 자동 마킹할 가능성 있음 (의도적)
- 한국어 unicode 자연 처리 — paragraph.text 가 str
- DOCX 는 page 개념이 없음 (워드 wrap 의존) → ExtractedSection.page = None

graceful degrade
- python-docx 가 corrupted DOCX 에서 raise → 본 파서가 RuntimeError 로 wrap (HwpxParser 패턴)
- 단락/표 단위 부분 실패는 warnings 누적
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import PurePosixPath

import docx as python_docx

from app.adapters.parser import ExtractedSection, ExtractionResult

logger = logging.getLogger(__name__)


# heading 판별 — paragraph.style.name 정규식
# python-docx 의 표준 heading 스타일은 "Heading 1" ~ "Heading 9" + "Title".
# 한국어 워드 사용자 정의 스타일도 "제목 1" 같은 패턴 → HwpxParser 의 정규식 재사용.
_HEADING_STYLE_PATTERN = re.compile(
    r"^(Heading\s*\d*|Title|Subtitle"
    r"|법-제목|제목|소제목|머리말|간지\d*|장|절|조|편|관|항목제목"
    r"|별표(\s*-.*)?|별첨(\s*-.*)?|개요\s*\d*|chapter\s*\d*)$",
    re.IGNORECASE,
)

# heading 판별 — inline 텍스트 패턴 fallback (HwpxParser 동일)
_HEADING_TEXT_PATTERN = re.compile(
    r"^(제\s*\d+\s*[조항장절편관]|부칙|별표\s*\d*|별첨\s*\d*)([\s(].*)?$"
)
_HEADING_TEXT_MAX_LEN = 80


class DocxParser:
    source_type = "docx"

    def can_parse(self, file_name: str, mime_type: str | None) -> bool:
        ext = PurePosixPath(file_name).suffix.lower()
        return ext == ".docx"

    def parse(self, data: bytes, *, file_name: str) -> ExtractionResult:
        sections: list[ExtractedSection] = []
        warnings: list[str] = []
        raw_parts: list[str] = []

        try:
            doc = python_docx.Document(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(
                f"DOCX 파서 초기화 실패: {file_name}: {exc}"
            ) from exc

        current_title: str | None = None

        # iter_inner_content (python-docx 1.x) — paragraph 와 table 을 XML 순서대로 반환
        try:
            inner = doc.iter_inner_content()
        except AttributeError:
            # 0.x fallback — paragraphs 우선 + tables 후순회 (순서 손실 trade-off)
            inner = list(doc.paragraphs) + list(doc.tables)

        for content in inner:
            try:
                if isinstance(content, python_docx.text.paragraph.Paragraph):
                    text = (content.text or "").strip()
                    if not text:
                        continue
                    style_name = (
                        getattr(content.style, "name", None) if content.style else None
                    )
                    if _is_heading_paragraph(text, style_name):
                        current_title = text
                    sections.append(
                        ExtractedSection(
                            text=text,
                            page=None,
                            section_title=current_title,
                            bbox=None,
                        )
                    )
                    raw_parts.append(text)
                elif isinstance(content, python_docx.table.Table):
                    table_text = _table_to_text(content)
                    if not table_text:
                        continue
                    sections.append(
                        ExtractedSection(
                            text=table_text,
                            page=None,
                            section_title=current_title,
                            bbox=None,
                        )
                    )
                    raw_parts.append(table_text)
            except Exception as exc:  # noqa: BLE001
                msg = f"DOCX 단락/표 추출 실패: {exc}"
                warnings.append(msg)
                logger.warning("%s (file=%s)", msg, file_name)
                continue

        return ExtractionResult(
            source_type=self.source_type,
            sections=sections,
            raw_text="\n\n".join(raw_parts),
            warnings=warnings,
        )


def _is_heading_paragraph(text: str, style_name: str | None) -> bool:
    """단락이 heading 인지 판정 (HwpxParser 와 동일 알고리즘).

    (A) Style.name 이 heading 정규식과 매칭되면 True.
    (B) 그렇지 않더라도 텍스트가 inline outline 패턴이고 길이 ≤ `_HEADING_TEXT_MAX_LEN`
        이면 True (본문 prefix false positive 차단).
    """
    if style_name and _HEADING_STYLE_PATTERN.match(style_name.strip()):
        return True
    if len(text) <= _HEADING_TEXT_MAX_LEN and _HEADING_TEXT_PATTERN.match(text):
        return True
    return False


def _table_to_text(table) -> str:
    """표를 chunk_filter table_noise 룰이 마킹 가능한 형태로 텍스트화.

    각 행은 `cell1 | cell2 | ... | cellN` 으로, 행 사이는 `\\n`. 빈 셀은 공란.
    빈 표는 빈 문자열 반환.
    """
    rows_text: list[str] = []
    for row in table.rows:
        try:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        except Exception:  # noqa: BLE001 — 행 단위 실패 허용
            continue
    return "\n".join(rows_text)
